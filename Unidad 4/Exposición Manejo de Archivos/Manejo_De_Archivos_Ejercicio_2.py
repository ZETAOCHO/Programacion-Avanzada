# Escribe un programa en Python que:
# 1.Cree una carpeta llamada Mensajes
# 2.Dentro de esa carpeta, cree un archivo de Word llamado bienvenida.docx
# 3.El archivo debe contener un título “¡Bienvenido al curso!” y un mensaje que diga:
# “Este curso te ayudará a aprender Python fácilmente.”
import os
from docx import Document
ruta_actual = os.path.dirname(os.path.abspath(__file__))  # Obtener la ruta actual del directorio de trabajo
print(f"Ruta actual: {ruta_actual}")  # Imprimir la ruta actual para verificar
# Crear carpeta "Mensajes" si no existe
if not os.path.exists(os.path.join(ruta_actual, 'Exposición Manejo de Archivos Mensajes')):
    print("Creando la carpeta 'Exposición Manejo de Archivos Mensajes'...")
    os.mkdir(os.path.join(ruta_actual, 'Exposición Manejo de Archivos Mensajes'))
else:
    print("La carpeta 'Exposición Manejo de Archivos Mensajes' ya existe.")

# Crear documento de Word
doc = Document()
doc.add_heading("¡Bienvenido al curso!", level=1)
doc.add_paragraph("Este curso te ayudará a aprender Python fácilmente.")

# Guardar documento en la carpeta "Mensajes"
doc.save(os.path.join(ruta_actual, "Exposición Manejo de Archivos Mensajes/bienvenida.docx"))